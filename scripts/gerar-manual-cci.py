#!/usr/bin/env python3
"""Gera os downloads oficiais do Manual CCI a partir da fonte JSON do app."""

import argparse
import hashlib
import json
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / 'manual-cci-conteudo.json'
OUT_DIR = ROOT / 'downloads'
DOCX_OUT = OUT_DIR / 'Manual_Operacional_CCI.docx'
PDF_OUT = OUT_DIR / 'Manual_Operacional_CCI.pdf'
MANIFEST = OUT_DIR / 'manual-cci-manifest.json'

BLUE = '2454D7'
DARK_BLUE = '0B2545'
LIGHT_BLUE = 'E8EEF5'
PALE_BLUE = 'EFF6FF'
PALE_GOLD = 'FFF7ED'
MUTED = '64748B'


def load_content():
    return json.loads(SOURCE.read_text(encoding='utf-8'))


def source_hash():
    return hashlib.sha256(SOURCE.read_bytes()).hexdigest()


def file_hash(path):
    return hashlib.sha256(path.read_bytes()).hexdigest() if path.exists() else None


def data_extenso_br(valor):
    ano, mes, dia = str(valor).split('-')
    meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
             'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
    return f"{int(dia)} de {meses[int(mes) - 1]} de {ano}"


def update_manifest(content):
    manifest = {
        'manual_version': content['manual_version'],
        'updated_at': content['updated_at'],
        'source_sha256': source_hash(),
        'files': {}
    }
    for fmt, path in [('docx', DOCX_OUT), ('pdf', PDF_OUT)]:
        if path.exists():
            manifest['files'][fmt] = {
                'name': path.name,
                'sha256': file_hash(path),
                'size': path.stat().st_size
            }
    MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for tag, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn('w:' + tag))
        if node is None:
            node = OxmlElement('w:' + tag)
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn('w:cantSplit')) is None:
            tr_pr.append(OxmlElement('w:cantSplit'))
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in('w:tcW')
            tc_w.set(qn('w:w'), str(widths[idx]))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)


def add_numbering(document, ordered):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    numbering = document.part.numbering_part.element
    abstract_ids = [int(x.get(qn('w:abstractNumId'))) for x in numbering.findall(qn('w:abstractNum'))]
    num_ids = [int(x.get(qn('w:numId'))) for x in numbering.findall(qn('w:num'))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1
    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(abstract_id))
    multi = OxmlElement('w:multiLevelType'); multi.set(qn('w:val'), 'singleLevel'); abstract.append(multi)
    lvl = OxmlElement('w:lvl'); lvl.set(qn('w:ilvl'), '0'); abstract.append(lvl)
    start = OxmlElement('w:start'); start.set(qn('w:val'), '1'); lvl.append(start)
    num_fmt = OxmlElement('w:numFmt'); num_fmt.set(qn('w:val'), 'decimal' if ordered else 'bullet'); lvl.append(num_fmt)
    lvl_text = OxmlElement('w:lvlText'); lvl_text.set(qn('w:val'), '%1.' if ordered else '•'); lvl.append(lvl_text)
    suff = OxmlElement('w:suff'); suff.set(qn('w:val'), 'tab'); lvl.append(suff)
    p_pr = OxmlElement('w:pPr')
    tabs = OxmlElement('w:tabs'); tab = OxmlElement('w:tab'); tab.set(qn('w:val'), 'num'); tab.set(qn('w:pos'), '540'); tabs.append(tab); p_pr.append(tabs)
    ind = OxmlElement('w:ind'); ind.set(qn('w:left'), '540'); ind.set(qn('w:hanging'), '270'); p_pr.append(ind); lvl.append(p_pr)
    numbering.append(abstract)
    num = OxmlElement('w:num'); num.set(qn('w:numId'), str(num_id))
    abstract_ref = OxmlElement('w:abstractNumId'); abstract_ref.set(qn('w:val'), str(abstract_id)); num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_list_number(paragraph, num_id):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0'); num_pr.append(ilvl)
    num = OxmlElement('w:numId'); num.set(qn('w:val'), str(num_id)); num_pr.append(num)
    p_pr.append(num_pr)


def add_docx_block(document, block, bullet_num, ordered_num):
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    kind = block['type']
    if kind == 'paragraph':
        document.add_paragraph(block['text'])
    elif kind == 'callout':
        table = document.add_table(rows=1, cols=1)
        table.autofit = False
        set_table_geometry(table, [9360])
        cell = table.cell(0, 0)
        fill = PALE_GOLD if block.get('tone') == 'warning' else PALE_BLUE
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); cell._tc.get_or_add_tcPr().append(shd)
        p = cell.paragraphs[0]
        lead = p.add_run(block.get('label', '') + ': '); lead.bold = True; lead.font.color.rgb = RGBColor.from_string(BLUE if block.get('tone') != 'warning' else '9A3412')
        p.add_run(block.get('text', ''))
        document.add_paragraph().paragraph_format.space_after = Pt(0)
    elif kind in ('bullets', 'checklist', 'steps'):
        for item in block['items']:
            text = item
            if kind == 'checklist': text = '☐ ' + item
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            if kind != 'checklist': set_list_number(p, ordered_num if kind == 'steps' else bullet_num)
            else: p.paragraph_format.left_indent = Pt(13.5)
            p.add_run(text)
    elif kind == 'definitions':
        for item in block['items']:
            p = document.add_paragraph()
            p.add_run(item['term'] + ': ').bold = True
            p.add_run(item['definition'])
    elif kind == 'table':
        cols = len(block['headers'])
        table = document.add_table(rows=1, cols=cols)
        table.style = 'Table Grid'; table.autofit = False
        widths = [4680, 4680] if cols == 2 else [2100, 2460, 4800]
        set_table_geometry(table, widths)
        for idx, header in enumerate(block['headers']):
            cell = table.rows[0].cells[idx]; cell.text = header; cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), LIGHT_BLUE); cell._tc.get_or_add_tcPr().append(shd)
            for run in cell.paragraphs[0].runs: run.bold = True; run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        table.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
        for source_row in block['rows']:
            row = table.add_row()
            for idx, value in enumerate(source_row):
                row.cells[idx].text = str(value); row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row.cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row.cells[idx].paragraphs[0].paragraph_format.left_indent = Pt(0)
                row.cells[idx].paragraphs[0].paragraph_format.first_line_indent = Pt(0)
        set_table_geometry(table, widths)
        document.add_paragraph().paragraph_format.space_after = Pt(0)


def generate_docx(content):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = document.styles
    normal = styles['Normal']; normal.font.name = 'Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri'); normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for style_name, size, before, after, color in [('Heading 1', 16, 18, 10, BLUE), ('Heading 2', 13, 14, 7, BLUE), ('Heading 3', 12, 10, 5, DARK_BLUE)]:
        style = styles[style_name]; style.font.name = 'Calibri'; style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri'); style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color); style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after); style.paragraph_format.left_indent = Pt(0); style.paragraph_format.first_line_indent = Pt(0); style.paragraph_format.keep_together = True; style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]; header.text = 'CCI  |  MANUAL OPERACIONAL'; header.style = styles['Normal']; header.runs[0].font.size = Pt(8); header.runs[0].font.bold = True; header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run(f"Uso interno  •  v{content['manual_version']}  •  {content['updated_at'].split('-')[::-1][0]}/{content['updated_at'].split('-')[::-1][1]}/{content['updated_at'].split('-')[::-1][2]}").font.size = Pt(8)

    for _ in range(4): document.add_paragraph()
    kicker = document.add_paragraph(); kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = kicker.add_run('GUIA INTERNO DE OPERAÇÃO E TRANSIÇÃO'); run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor.from_string(BLUE)
    title = document.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; title.paragraph_format.space_after = Pt(10); run = title.add_run(content['title']); run.bold = True; run.font.size = Pt(28); run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    subtitle = document.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER; subtitle.paragraph_format.space_after = Pt(10); run = subtitle.add_run(content['subtitle']); run.font.size = Pt(14); run.font.color.rgb = RGBColor.from_string(BLUE)
    meta = document.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = meta.add_run(f"Versão {content['manual_version']}  |  {data_extenso_br(content['updated_at'])}  |  {content['classification']}"); run.font.size = Pt(10); run.font.color.rgb = RGBColor.from_string(MUTED)
    document.add_paragraph()
    add_docx_block(document, {'type': 'callout', 'tone': 'info', 'label': 'Situação atual', 'text': content['status']['summary']}, 0, 0)
    document.add_paragraph()
    brand = document.add_paragraph(); brand.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = brand.add_run('SP Assessoria Contábil  •  Consultor Contábil Inteligente'); run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor.from_string(MUTED)
    document.add_page_break()

    bullet_num = add_numbering(document, False); ordered_num = add_numbering(document, True)
    for chapter in content['chapters']:
        if chapter.get('page_break_before'):
            document.add_page_break()
        document.add_heading(chapter['title'], level=1)
        for block in chapter['blocks']:
            add_docx_block(document, block, bullet_num, ordered_num)
    document.core_properties.title = content['title']
    document.core_properties.subject = content['subtitle']
    document.core_properties.author = 'SP Assessoria Contábil'
    document.core_properties.comments = f"Fonte oficial: manual-cci-conteudo.json; sha256={source_hash()}"
    document.save(DOCX_OUT)


def generate_pdf(content):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, PageBreak, Paragraph, Spacer, Table, TableStyle, KeepTogether

    styles = getSampleStyleSheet()
    body = ParagraphStyle('BodyCCI', parent=styles['BodyText'], fontName='Helvetica', fontSize=9.5, leading=12, spaceAfter=6, textColor=colors.HexColor('#0F172A'))
    h1 = ParagraphStyle('H1CCI', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=16, leading=19, spaceBefore=15, spaceAfter=9, textColor=colors.HexColor('#2454D7'))
    title = ParagraphStyle('TitleCCI', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=28, leading=33, alignment=TA_CENTER, textColor=colors.HexColor('#0B2545'), spaceAfter=10)
    subtitle = ParagraphStyle('SubtitleCCI', parent=body, fontSize=13, leading=17, alignment=TA_CENTER, textColor=colors.HexColor('#2454D7'), spaceAfter=10)
    small = ParagraphStyle('SmallCCI', parent=body, fontSize=8, leading=10, textColor=colors.HexColor('#64748B'))
    small_center = ParagraphStyle('SmallCenterCCI', parent=small, alignment=TA_CENTER)
    label = ParagraphStyle('LabelCCI', parent=body, fontName='Helvetica-Bold', textColor=colors.HexColor('#153A67'), spaceAfter=1)
    list_style = ParagraphStyle('ListCCI', parent=body, leftIndent=16, firstLineIndent=-12, spaceAfter=4)

    def page(canvas, doc):
        canvas.saveState(); canvas.setFont('Helvetica-Bold', 7); canvas.setFillColor(colors.HexColor('#64748B')); canvas.drawString(inch, letter[1] - 0.55 * inch, 'CCI  |  MANUAL OPERACIONAL')
        canvas.setFont('Helvetica', 7); canvas.drawRightString(letter[0] - inch, 0.5 * inch, f"Uso interno  •  v{content['manual_version']}  •  Página {doc.page}"); canvas.restoreState()

    document = BaseDocTemplate(str(PDF_OUT), pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=0.8 * inch, bottomMargin=0.75 * inch, title=content['title'], author='SP Assessoria Contábil')
    frame = Frame(document.leftMargin, document.bottomMargin, document.width, document.height, id='normal')
    document.addPageTemplates([PageTemplate(id='cci', frames=frame, onPage=page)])
    story = [Spacer(1, 1.25 * inch), Paragraph('GUIA INTERNO DE OPERAÇÃO E TRANSIÇÃO', ParagraphStyle('Kicker', parent=small_center, fontName='Helvetica-Bold', textColor=colors.HexColor('#2454D7'), spaceAfter=14)), Paragraph(content['title'], title), Paragraph(content['subtitle'], subtitle), Paragraph(f"Versão {content['manual_version']}  |  {data_extenso_br(content['updated_at'])}  |  {content['classification']}", small_center), Spacer(1, 0.35 * inch)]
    status_table = Table([[Paragraph('<b>Situação atual:</b> ' + content['status']['summary'], body)]], colWidths=[document.width])
    status_table.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#EFF6FF')), ('BOX', (0,0), (-1,-1), 0, colors.white), ('LINEBEFORE', (0,0), (0,-1), 3, colors.HexColor('#2454D7')), ('LEFTPADDING', (0,0), (-1,-1), 10), ('RIGHTPADDING', (0,0), (-1,-1), 10), ('TOPPADDING', (0,0), (-1,-1), 9), ('BOTTOMPADDING', (0,0), (-1,-1), 9)]))
    story += [status_table, Spacer(1, 0.55 * inch), Paragraph('SP Assessoria Contábil  •  Consultor Contábil Inteligente', small_center), PageBreak()]

    for chapter in content['chapters']:
        if chapter.get('page_break_before'):
            story.append(PageBreak())
        story.append(Paragraph(chapter['title'], h1))
        for block in chapter['blocks']:
            kind = block['type']
            if kind == 'paragraph': story.append(Paragraph(block['text'], body))
            elif kind == 'callout':
                fill = '#FFF7ED' if block.get('tone') == 'warning' else '#EFF6FF'; line = '#F59E0B' if block.get('tone') == 'warning' else '#2454D7'
                table = Table([[Paragraph(f"<b>{block.get('label','')}:</b> {block.get('text','')}", body)]], colWidths=[document.width])
                table.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,-1),colors.HexColor(fill)),('LINEBEFORE',(0,0),(0,-1),3,colors.HexColor(line)),('LEFTPADDING',(0,0),(-1,-1),9),('RIGHTPADDING',(0,0),(-1,-1),9),('TOPPADDING',(0,0),(-1,-1),7),('BOTTOMPADDING',(0,0),(-1,-1),7)])); story += [table, Spacer(1, 5)]
            elif kind == 'definitions':
                for item in block['items']: story.append(KeepTogether([Paragraph(item['term'], label), Paragraph(item['definition'], body)]))
            elif kind in ('bullets', 'checklist', 'steps'):
                for index, item in enumerate(block['items'], 1):
                    marker = f'{index}.' if kind == 'steps' else '[ ]' if kind == 'checklist' else '•'
                    story.append(Paragraph(f'{marker} {item}', list_style))
                story.append(Spacer(1, 4))
            elif kind == 'table':
                cols = len(block['headers']); widths = [document.width/2]*2 if cols == 2 else [document.width*0.22, document.width*0.25, document.width*0.53]
                rows = [[Paragraph(str(cell), label) for cell in block['headers']]] + [[Paragraph(str(cell), small) for cell in row] for row in block['rows']]
                table = Table(rows, colWidths=widths, repeatRows=1, hAlign='LEFT')
                table.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#E8EEF5')),('GRID',(0,0),(-1,-1),0.5,colors.HexColor('#94A3B8')),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LEFTPADDING',(0,0),(-1,-1),6),('RIGHTPADDING',(0,0),(-1,-1),6),('TOPPADDING',(0,0),(-1,-1),6),('BOTTOMPADDING',(0,0),(-1,-1),6)])); story += [table, Spacer(1, 5)]
    document.build(story)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--format', choices=['docx', 'pdf', 'all'], default='all')
    args = parser.parse_args()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    content = load_content()
    if args.format in ('docx', 'all'): generate_docx(content)
    if args.format in ('pdf', 'all'): generate_pdf(content)
    update_manifest(content)
    print(DOCX_OUT if args.format == 'docx' else PDF_OUT if args.format == 'pdf' else OUT_DIR)


if __name__ == '__main__':
    main()
